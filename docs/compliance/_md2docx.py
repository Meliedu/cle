"""Minimal Markdown -> .docx converter for the HKUST compliance doc.

Handles: ATX headings, bullet lists, ordered lists, blockquotes, GFM tables,
inline **bold**, *italic*, `code`. Good enough for this one document; not a
general-purpose converter. Run with the backend venv python.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, DST = sys.argv[1], sys.argv[2]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def add_runs(paragraph, text):
    text = text.replace("\\|", "|")
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def split_row(line):
    line = line.strip().strip("|")
    # split on unescaped pipes
    cells, buf, i = [], "", 0
    while i < len(line):
        if line[i] == "\\" and i + 1 < len(line):
            buf += line[i:i+2]; i += 2; continue
        if line[i] == "|":
            cells.append(buf.strip()); buf = ""; i += 1; continue
        buf += line[i]; i += 1
    cells.append(buf.strip())
    return cells


def emit_table(rows):
    header, *body = rows
    cols = split_row(header)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], c)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for line in body:
        if re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line):  # separator row
            continue
        vals = split_row(line)
        cells = t.add_row().cells
        for j in range(len(cols)):
            cells[j].paragraphs[0].text = ""
            add_runs(cells[j].paragraphs[0], vals[j] if j < len(vals) else "")
            for run in cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)


lines = open(SRC, encoding="utf-8").read().splitlines()
i = 0
while i < len(lines):
    line = lines[i]

    # table block
    if line.lstrip().startswith("|") and i + 1 < len(lines) and re.search(r"\|[\s:|-]+\|", lines[i+1]):
        block = []
        while i < len(lines) and lines[i].lstrip().startswith("|"):
            block.append(lines[i]); i += 1
        emit_table(block)
        continue

    stripped = line.strip()

    if not stripped:
        i += 1; continue

    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        h = doc.add_heading(level=min(level, 4))
        add_runs(h, m.group(2))
        i += 1; continue

    if stripped.startswith(">"):
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, stripped.lstrip(">").strip())
        i += 1; continue

    if stripped in ("---", "***", "___"):
        doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue

    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, m.group(2))
        i += 1; continue

    if stripped.startswith(("- ", "* ")):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, stripped[2:])
        i += 1; continue

    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(DST)
print("wrote", DST)
