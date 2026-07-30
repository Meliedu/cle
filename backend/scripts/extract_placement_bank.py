"""Extract the Meli placement item bank from the v1.2 teacher re-review package.

Why this exists
---------------
v1.1 shipped ``data/selected_items_v1.1.json``. **v1.2 does not ship any
machine-readable item file** -- the package is PDFs, DOCX sources and one XLSX.
So the bank has to be reconstructed, and the reconstruction has to be
reproducible and checkable rather than a one-off paste.

The two halves come from two different files on purpose, and that split is the
whole security story of this test:

* the **student-safe** surface (passage, stem, options) is read from the student
  booklet ``editable/Meli_Placement_Test_v1.2_Form_?.docx`` -- the same artefact
  a learner receives, so nothing restricted can leak in by construction;
* the **restricted** surface (key, answer text, rationale, listening transcript,
  reference band) is read from the controlled workbook + Administrator Pack.

Reading the safe fields from the student booklet is not a convenience. If we
derived them from the Administrator Pack we would have no structural guarantee
that a key or a band label never rides along into the delivery payload.

Usage::

    python scripts/extract_placement_bank.py --package <dir> [--out <file>]

Output is a single versioned JSON consumed by
``app/services/placement_bank.py``. Nothing here runs at request time.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import docx  # type: ignore[import-untyped]
import openpyxl

VERSION_CODE = "meli-placement-v1.2-candidate-2026-07-28"
SOURCE_PACKAGE = "Meli_Placement_Test_v1.2_Teacher_Rereview_2026-07-28"

FORMS = ("A", "B", "C", "D", "E")
ITEMS_PER_FORM = 30
BANDS = (1, 2, 3, 4, 5, 6)

#: Section boundaries are fixed by the published blueprint (12 / 6 / 12) and are
#: asserted rather than inferred, so a re-ordered source fails loudly.
SECTION_BY_QUESTION = {
    **{q: "listening" for q in range(1, 13)},
    **{q: "language_use" for q in range(13, 19)},
    **{q: "reading" for q in range(19, 31)},
}

#: Workbook ``Skill`` -> our section key. The workbook says "Vocabulary" for the
#: middle block while every learner-facing artefact says "Language Use"; the
#: booklet heading and the Summary sheet both use "Language Use", so that wins.
SKILL_TO_SECTION = {
    "Listening": "listening",
    "Vocabulary": "language_use",
    "Language Use": "language_use",
    "Reading": "reading",
}

OPTION_LETTERS = ("A", "B", "C", "D", "E", "F")

#: A line like "A. 老师     B. 医生" -- the booklet lays options two per line.
_OPTION_PAIR = re.compile(
    r"([A-F])\.\s*(.*?)(?=\s{2,}[A-F]\.|$)",
    re.DOTALL,
)
_QUESTION_START = re.compile(r"^(\d{1,2})\.\s*(.*)$", re.DOTALL)
_ADMIN_HEADING = re.compile(r"^Q(\d{1,2})\s*·\s*(\S+)\s*·\s*legacy HSK(\d)$")
_ADMIN_KEY = re.compile(r"^Key:\s*([A-F](?:-[A-F])*)\s*·\s*(.*?)(?:\s*\|\s*(.*))?$")
_CLOZE_PROMPT = re.compile(r"^请选择最合适的一项填入第（(\d{1,2})）空。$")
_CLOZE_BLANK = re.compile(r"（(\d{1,2})）\s*_")
_ORDER_LINE = "Order / 顺序"

#: Booklet furniture that sits between questions and must never be mistaken for
#: item text. Without this, a section heading and its "Questions 19-30 · about
#: 13 minutes" descriptor get absorbed as the stem of the *preceding* item.
_FURNITURE = (
    re.compile(r"^(Listening|Language Use|Reading|Instructions|Answer Sheet)\s*/"),
    re.compile(r"^Questions \d{1,2}-\d{1,2}\s*·"),
    re.compile(r"^Internal use\s*·"),
    re.compile(r"^MELI\s*·"),
    re.compile(r"^Meli\s*·\s*CLE"),
    re.compile(r"^Form [A-E]\s*·"),
    re.compile(r"^30-minute internal diagnostic"),
    re.compile(r"^Do not open the next page"),
)


def _is_furniture(line: str) -> bool:
    return any(pattern.match(line) for pattern in _FURNITURE)


class ExtractError(RuntimeError):
    """Raised when the source package does not match the published blueprint."""


@dataclass
class RawItem:
    """One question as read from the student booklet (safe fields only)."""

    question_number: int
    section: str
    passage: str | None = None
    stem: str = ""
    prompt: str | None = None
    options: list[dict[str, str]] = field(default_factory=list)
    is_sequence: bool = False


def _clean(text: str) -> str:
    """Normalise whitespace without touching CJK content.

    NFC only. The booklets mix ideographic space (U+3000) inside cloze brackets
    with ASCII spaces around option letters; collapsing the former would corrupt
    the rendered blank, so only ASCII runs are collapsed.
    """
    normalised = unicodedata.normalize("NFC", text)
    # Word emits NBSP inside option runs; fold it to a plain space so option
    # text compares cleanly. U+3000 is deliberately NOT folded: it is the
    # rendered blank inside a cloze bracket, so losing it changes the item.
    normalised = normalised.replace(" ", " ")
    return re.sub(r"[ \t]+", " ", normalised).strip()


def _parse_options(line: str) -> list[tuple[str, str]]:
    """Pull ``(letter, text)`` pairs out of one booklet option line."""
    found = [
        (m.group(1), _clean(m.group(2)))
        for m in _OPTION_PAIR.finditer(line)
    ]
    return [(letter, text) for letter, text in found if text]


def _looks_like_options(line: str) -> bool:
    return bool(re.match(r"^[A-F]\.\s", line.strip()))


def read_student_form(path: Path, form_code: str) -> list[RawItem]:
    """Read the 30 safe question surfaces out of one student booklet."""
    document = docx.Document(str(path))
    # Raw text is kept alongside the cleaned text on purpose: the booklet marks
    # an option boundary with a run of spaces ("A. 老师     B. 医生"), and
    # collapsing that run first would destroy the only separator there is.
    paragraphs = [(p.text, _clean(p.text)) for p in document.paragraphs]

    items: list[RawItem] = []
    current: RawItem | None = None
    # Text lines accumulated since the question opened but before its options:
    # the first is the stimulus, the last before options is the actual question.
    body: list[str] = []

    def close() -> None:
        nonlocal current, body
        if current is None:
            return
        _finalise(current, body)
        items.append(current)
        current = None
        body = []

    for raw_line, line in paragraphs:
        if not line:
            continue
        if line.startswith(_ORDER_LINE):
            if current is not None:
                current.is_sequence = True
            continue
        if line.startswith("Answer Sheet"):
            break
        if _is_furniture(line):
            continue

        match = _QUESTION_START.match(line)
        # A numbered line only opens a question when the number is exactly the
        # next one we expect. Passages legitimately contain "…3.5倍" style text,
        # and option lines start with a letter, so both are rejected here.
        expected_next = (current.question_number + 1) if current else 1
        if (
            match
            and int(match.group(1)) == expected_next
            and not _looks_like_options(line)
        ):
            close()
            number = int(match.group(1))
            current = RawItem(
                question_number=number,
                section=SECTION_BY_QUESTION[number],
            )
            body = [_clean(match.group(2))] if match.group(2).strip() else []
            continue

        if current is None:
            continue

        if _looks_like_options(line):
            for letter, text in _parse_options(raw_line):
                current.options.append({"letter": letter, "text": text})
            continue

        cloze = _CLOZE_PROMPT.match(line)
        if cloze:
            current.prompt = line
            continue

        body.append(line)

    close()

    if len(items) != ITEMS_PER_FORM:
        raise ExtractError(
            f"Form {form_code}: expected {ITEMS_PER_FORM} questions, parsed {len(items)}"
        )
    return items


def _finalise(item: RawItem, body: list[str]) -> None:
    """Split the accumulated text into passage vs stem.

    A single line is the stem. Two or more means the last is the question and
    everything before it is the stimulus the question is asked about.
    """
    lines = [line for line in body if line]
    if not lines:
        item.stem = ""
    elif len(lines) == 1:
        item.stem = lines[0]
    else:
        item.passage = "\n".join(lines[:-1])
        item.stem = lines[-1]

    # Cloze items carry their instruction in `prompt`; the passage IS the item,
    # so a cloze whose text collapsed into `stem` should present as a passage.
    if item.prompt and item.passage is None and item.stem:
        item.passage = item.stem
        item.stem = ""


def read_workbook_metadata(path: Path) -> dict[tuple[str, int], dict[str, Any]]:
    """Read the restricted per-item metadata from the scoring workbook."""
    workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    try:
        sheet = workbook["Item Metadata"]
        rows = list(sheet.iter_rows(values_only=True))
    finally:
        workbook.close()

    header_row = next(
        (i for i, row in enumerate(rows) if row and row[0] == "Form"), None
    )
    if header_row is None:
        raise ExtractError("Item Metadata sheet has no 'Form' header row")
    header = [str(c).strip() if c is not None else "" for c in rows[header_row]]
    index = {name: i for i, name in enumerate(header) if name}

    required = [
        "Form", "Q", "Item ID", "Slot", "Legacy band", "Skill", "Item type",
        "Response", "Topic", "Subskill", "Expected sec", "Playback", "Key",
    ]
    missing = [name for name in required if name not in index]
    if missing:
        raise ExtractError(f"Item Metadata is missing columns: {missing}")

    out: dict[tuple[str, int], dict[str, Any]] = {}
    for row in rows[header_row + 1:]:
        if not row or row[index["Form"]] is None:
            continue
        form = str(row[index["Form"]]).strip()
        number = int(row[index["Q"]])
        out[(form, number)] = {
            "item_id": str(row[index["Item ID"]]).strip(),
            "slot": str(row[index["Slot"]]).strip(),
            "legacy_band": int(row[index["Legacy band"]]),
            "skill": str(row[index["Skill"]]).strip(),
            "item_type": str(row[index["Item type"]]).strip(),
            "response_format": str(row[index["Response"]]).strip(),
            "topic": _optional(row, index, "Topic"),
            "subskill": _optional(row, index, "Subskill"),
            "expected_seconds": int(row[index["Expected sec"]]),
            "audio_playback": (
                int(row[index["Playback"]]) if row[index["Playback"]] else None
            ),
            "correct_answer": str(row[index["Key"]]).strip(),
            "target_vocabulary": _optional(row, index, "Target vocab"),
            "target_grammar": _optional(row, index, "Target grammar/discourse"),
            "copyright_status": _optional(row, index, "Copyright status"),
            "qa_status": _optional(row, index, "QA status"),
        }
    if len(out) != len(FORMS) * ITEMS_PER_FORM:
        raise ExtractError(
            f"Item Metadata: expected 150 rows, read {len(out)}"
        )
    return out


def _optional(row: tuple, index: dict[str, int], name: str) -> str | None:
    position = index.get(name)
    if position is None or position >= len(row):
        return None
    value = row[position]
    return str(value).strip() if value not in (None, "") else None


def read_admin_pack(path: Path) -> dict[tuple[str, int], dict[str, Any]]:
    """Read transcripts, answer texts and rationales from the Administrator Pack."""
    document = docx.Document(str(path))
    entries: dict[tuple[str, int], dict[str, Any]] = {}

    form_code: str | None = None
    current: tuple[str, int] | None = None
    buffer: list[str] = []

    def flush() -> None:
        if current is None:
            return
        entry = entries[current]
        # Position relative to the "Key:" line is what distinguishes a listening
        # script from a rationale: both are free prose. Before the key it is the
        # script the proctor reads; after it, the defence of the key. Matching on
        # a prefix instead would silently drop every rationale that opens with
        # "The decisive evidence is…" rather than "Option D…".
        transcript: list[str] = []
        rationale: list[str] = []
        seen_key = False
        for line in buffer:
            if line.startswith("Key:"):
                seen_key = True
                match = _ADMIN_KEY.match(line)
                if match:
                    entry["answer_text"] = _clean(match.group(2))
                    if match.group(3):
                        rationale.append(_clean(match.group(3)))
            elif line.startswith("Question:"):
                entry["admin_question"] = _clean(line[len("Question:"):])
            elif line.startswith("Target:"):
                entry["admin_target"] = _clean(line[len("Target:"):])
            elif seen_key:
                rationale.append(line)
            else:
                transcript.append(line)
        if transcript:
            entry["transcript"] = "\n".join(transcript)
        if rationale:
            entry["rationale"] = " ".join(rationale)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name

        if style == "Heading 1":
            match = re.match(r"^Form ([A-E]):", text)
            if match:
                flush()
                current, buffer = None, []
                form_code = match.group(1)
            continue

        if style == "Heading 3" and form_code:
            heading = _ADMIN_HEADING.match(text)
            if heading:
                flush()
                number = int(heading.group(1))
                current = (form_code, number)
                entries[current] = {
                    "item_id": heading.group(2),
                    "legacy_band": int(heading.group(3)),
                }
                buffer = []
            continue

        if current is not None:
            # python-docx collapses the soft line breaks inside a script cell
            # into a single paragraph; keep both speaker turns.
            buffer.extend(_clean(part) for part in paragraph.text.split("\n") if part.strip())

    flush()
    return entries


def build_bank(package: Path, teacher_flags: dict[str, list[dict[str, str]]]) -> dict[str, Any]:
    """Assemble and validate the full versioned bank."""
    root = _resolve_root(package)
    editable = root / "editable"
    admin = root / "admin"

    workbook_path = admin / "Meli_Placement_Test_v1.2_Scoring_and_Calibration.xlsx"
    admin_pack_path = editable / "Meli_Placement_Test_v1.2_Administrator_Pack.docx"
    for path in (workbook_path, admin_pack_path):
        if not path.exists():
            raise ExtractError(f"missing source file: {path}")

    metadata = read_workbook_metadata(workbook_path)
    admin_entries = read_admin_pack(admin_pack_path)

    forms: list[dict[str, Any]] = []
    checksums: dict[str, str] = {
        workbook_path.name: _sha256(workbook_path),
        admin_pack_path.name: _sha256(admin_pack_path),
    }

    for form_code in FORMS:
        booklet = editable / f"Meli_Placement_Test_v1.2_Form_{form_code}.docx"
        if not booklet.exists():
            raise ExtractError(f"missing student booklet: {booklet}")
        checksums[booklet.name] = _sha256(booklet)

        raw_items = read_student_form(booklet, form_code)
        items: list[dict[str, Any]] = []

        for raw in raw_items:
            meta = metadata.get((form_code, raw.question_number))
            if meta is None:
                raise ExtractError(
                    f"Form {form_code} Q{raw.question_number}: no workbook metadata"
                )
            admin_entry = admin_entries.get((form_code, raw.question_number), {})

            if admin_entry and admin_entry.get("item_id") != meta["item_id"]:
                raise ExtractError(
                    f"Form {form_code} Q{raw.question_number}: item id mismatch "
                    f"({admin_entry.get('item_id')} vs {meta['item_id']})"
                )

            expected_section = SKILL_TO_SECTION.get(meta["skill"])
            if expected_section != raw.section:
                raise ExtractError(
                    f"Form {form_code} Q{raw.question_number}: skill "
                    f"{meta['skill']!r} does not match blueprint section {raw.section!r}"
                )

            items.append(
                _assemble_item(form_code, raw, meta, admin_entry, teacher_flags)
            )

        _validate_form(form_code, items)
        forms.append({"form_code": form_code, "items": items})

    return {
        "version_code": VERSION_CODE,
        "source_package": SOURCE_PACKAGE,
        "source_checksums": checksums,
        "blueprint": {
            "items_per_form": ITEMS_PER_FORM,
            "sections": {"listening": 12, "language_use": 6, "reading": 12},
            "items_per_band": 5,
            "bands": list(BANDS),
            "expected_seconds_per_form": 1401,
            "duration_minutes": 30,
        },
        "claim_boundary": {
            "purpose": (
                "This is an internal Meli/CLE placement screener. It is not an "
                "official HSK examination or score report."
            ),
            "privacy": (
                "Results support a course recommendation and CLE review. They do "
                "not create an official proficiency credential."
            ),
        },
        "forms": forms,
    }


def _assemble_item(
    form_code: str,
    raw: RawItem,
    meta: dict[str, Any],
    admin_entry: dict[str, Any],
    teacher_flags: dict[str, list[dict[str, str]]],
) -> dict[str, Any]:
    key = meta["correct_answer"]
    letters = [option["letter"] for option in raw.options]
    flags = list(teacher_flags.get(meta["item_id"], []))

    if raw.is_sequence:
        response_format = "sequence"
        valid = all(part in letters for part in key.split("-"))
    else:
        response_format = f"choice_{len(letters)}"
        valid = key in letters

    if not valid:
        # A key naming an option the booklet does not offer is unscoreable.
        # Recorded as a blocking flag rather than aborting: the bank must still
        # load so the review queue can show CLE exactly what is wrong.
        flags.append(
            {
                "code": "key_not_in_options",
                "detail": f"key {key!r} is not among the booklet options {letters}",
                "source": "extractor",
            }
        )

    flags.extend(_cloze_numbering_flags(raw))

    return {
        "question_number": raw.question_number,
        "section": raw.section,
        "response_format": response_format,
        "option_letters": letters,
        "expected_seconds": meta["expected_seconds"],
        "audio_playback": meta["audio_playback"],
        # --- student-safe delivery surface (from the student booklet only) ---
        "safe": {
            "passage": raw.passage,
            "stem": raw.stem,
            "prompt": raw.prompt,
            "options": raw.options,
        },
        # --- restricted: never serialised into a student response ---
        "restricted": {
            "item_id": meta["item_id"],
            "slot": meta["slot"],
            "legacy_band": meta["legacy_band"],
            "skill": meta["skill"],
            "item_type": meta["item_type"],
            "workbook_response_format": meta["response_format"],
            "topic": meta["topic"],
            "subskill": meta["subskill"],
            "correct_answer": key,
            "answer_text": admin_entry.get("answer_text"),
            "rationale": admin_entry.get("rationale"),
            "transcript": admin_entry.get("transcript"),
            "target_vocabulary": meta["target_vocabulary"],
            "target_grammar": meta["target_grammar"],
            "copyright_status": meta["copyright_status"],
            "qa_status": meta["qa_status"],
        },
        "teacher_flags": flags,
    }


def _cloze_numbering_flags(raw: RawItem) -> list[dict[str, str]]:
    """Flag a cloze whose blank is numbered differently from its question.

    The booklet writes the blank as "（17）____" and instructs "填入第（17）空".
    Both numbers must equal the question number, or the learner is told to fill
    a blank that is not on the page. This is a mechanical check, and it is what
    catches Form B Q17 without anyone having to hardcode that item.
    """
    flags: list[dict[str, str]] = []
    expected = raw.question_number

    for source, text in (("passage", raw.passage), ("prompt", raw.prompt)):
        if not text:
            continue
        pattern = _CLOZE_BLANK if source == "passage" else _CLOZE_PROMPT
        for match in pattern.finditer(text):
            found = int(match.group(1))
            if found != expected:
                flags.append(
                    {
                        "code": "cloze_blank_number_mismatch",
                        "detail": (
                            f"{source} numbers the blank （{found}） but this is "
                            f"question {expected}"
                        ),
                        "source": "extractor",
                    }
                )
    return flags


def _validate_form(form_code: str, items: list[dict[str, Any]]) -> None:
    """Assert the published blueprint on the reconstructed form."""
    if len(items) != ITEMS_PER_FORM:
        raise ExtractError(f"Form {form_code}: {len(items)} items")

    sections: dict[str, int] = {}
    bands: dict[int, int] = {}
    for item in items:
        sections[item["section"]] = sections.get(item["section"], 0) + 1
        band = item["restricted"]["legacy_band"]
        bands[band] = bands.get(band, 0) + 1

    if sections != {"listening": 12, "language_use": 6, "reading": 12}:
        raise ExtractError(f"Form {form_code}: section counts {sections}")
    if any(bands.get(band) != 5 for band in BANDS):
        raise ExtractError(f"Form {form_code}: band counts {bands}")

    total = sum(item["expected_seconds"] for item in items)
    if total != 1401:
        raise ExtractError(f"Form {form_code}: expected 1401 seconds, got {total}")

    for item in items:
        options = item["safe"]["options"]
        if len(options) < 2:
            raise ExtractError(
                f"Form {form_code} Q{item['question_number']}: {len(options)} options"
            )
        texts = [option["text"] for option in options]
        if len(set(texts)) != len(texts):
            raise ExtractError(
                f"Form {form_code} Q{item['question_number']}: duplicate option text"
            )
        if not item["safe"]["stem"] and not item["safe"]["passage"]:
            raise ExtractError(
                f"Form {form_code} Q{item['question_number']}: no stem or passage"
            )


def _resolve_root(package: Path) -> Path:
    """Find the directory that actually holds ``admin/`` and ``editable/``.

    The distributed zip unpacks with its own name repeated, so the path the
    user hands us may be the parent, the package dir, or the doubled dir.
    """
    candidates = [
        package,
        package / SOURCE_PACKAGE,
        package / SOURCE_PACKAGE / SOURCE_PACKAGE,
    ]
    for candidate in candidates:
        if (candidate / "admin").is_dir() and (candidate / "editable").is_dir():
            return candidate
    raise ExtractError(
        f"no admin/ + editable/ pair under {package} (tried: "
        + ", ".join(str(c) for c in candidates)
        + ")"
    )


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    digest.update(path.read_bytes())
    return digest.hexdigest()


#: Teacher annotations on the v1.2 PDFs (author ``lcxfxu``, 30 July 2026) that
#: name a defect in the **key**, keyed by item ID.
#:
#: The v1.2 comment PDFs carry ~59 markup annotations. Most are prose edits
#: (word choice, tightening) which are content work for CLE and do not change
#: what a response scores as. The three below do: two dispute the key outright,
#: and one disputes the wording of the keyed option itself. An item carrying a
#: ``key_disputed`` flag cannot be scored defensibly, so publication is blocked
#: until CLE resolves it (see ``services/placement_preflight.py``).
#:
#: Comments are recorded verbatim in the teacher's own words. Meli does not
#: adjudicate assessment content; it surfaces the dispute and refuses to
#: pretend the item is scoreable.
TEACHER_KEY_FLAGS: dict[str, list[dict[str, str]]] = {
    # Form B Q17. The blank-number defect is caught mechanically as well; this
    # records the second half of the same annotation, that the key is wrong.
    "HSK5-B-13": [
        {
            "code": "key_disputed",
            "detail": "没有正确选项，administration里面的答案不正确，修改时一起修改",
            "source": "teacher:lcxfxu:2026-07-30:Form_B_p4",
        }
    ],
    # Form C Q23. Two options are defensible, so the item has no single key.
    "HSK3-C-12": [
        {
            "code": "key_disputed",
            "detail": "选项A和C都可以",
            "source": "teacher:lcxfxu:2026-07-30:Form_C_p5",
        }
    ],
    # Form E Q17. The keyed option's wording is disputed, not the choice of
    # option, so this is a revision request rather than an unscoreable item.
    "HSK5-E-11": [
        {
            "code": "keyed_option_wording",
            "detail": "A:改成“隐瞒”会更好一些",
            "source": "teacher:lcxfxu:2026-07-30:Form_E_p4",
        }
    ],
}


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--package", required=True, type=Path)
    parser.add_argument(
        "--out",
        type=Path,
        default=Path(__file__).resolve().parents[1]
        / "app" / "data" / "placement" / "meli-placement-v1.2.json",
    )
    args = parser.parse_args(argv)

    bank = build_bank(args.package, TEACHER_KEY_FLAGS)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(
        json.dumps(bank, ensure_ascii=False, indent=1) + "\n", encoding="utf-8"
    )

    total = sum(len(form["items"]) for form in bank["forms"])
    flagged = sum(
        1
        for form in bank["forms"]
        for item in form["items"]
        if item["teacher_flags"]
    )
    print(f"wrote {args.out} :: {len(bank['forms'])} forms, {total} items, {flagged} flagged")
    return 0


if __name__ == "__main__":
    sys.exit(main())
